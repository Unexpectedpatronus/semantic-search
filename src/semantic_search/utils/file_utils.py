"""Утилиты для работы с файлами (рефакторинг)"""

from collections import Counter
from pathlib import Path
from typing import Generator, List, Optional

import pymupdf
from docx import Document as DocxDocument
from loguru import logger

from semantic_search.config import SUPPORTED_EXTENSIONS

try:
    import win32com.client

    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False
    logger.warning("pywin32 не установлен. Поддержка .doc файлов недоступна")


class FileExtractor:
    """Класс для извлечения текста из различных форматов файлов"""

    # Константы для ограничений
    MAX_PDF_PAGES_IN_MEMORY = 100
    PAGE_BATCH_SIZE = 10
    MIN_PAGE_TEXT_LENGTH = 50

    def __init__(self):
        self.word_app: Optional[object] = None
        self._init_word_app()

    def _init_word_app(self):
        """Инициализация Word Application"""
        if DOC_SUPPORT:
            try:
                self.word_app = win32com.client.Dispatch("Word.Application")
                self.word_app.Visible = False
            except Exception as e:
                logger.warning(f"Не удалось инициализировать Word Application: {e}")
                self.word_app = None

    def __del__(self):
        """Освобождение ресурсов"""
        if self.word_app is not None:
            try:
                if hasattr(self.word_app, "Quit"):
                    self.word_app.Quit()
            except Exception:
                pass

    def find_documents(self, root_path: Path) -> List[Path]:
        """
        Рекурсивный поиск документов в директории

        Args:
            root_path: Путь к корневой директории

        Returns:
            Список путей к найденным файлам
        """
        if not root_path.exists():
            raise FileNotFoundError(f"Директория не найдена: {root_path}")

        found_files = []
        logger.info(f"Поиск документов в: {root_path}")

        for file_path in root_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                found_files.append(file_path)

        # Логирование статистики
        self._log_file_statistics(found_files)

        return found_files

    def _log_file_statistics(self, files: List[Path]) -> None:
        """Логирование статистики найденных файлов"""
        logger.info(f"Найдено файлов: {len(files)}")
        ext_counter = Counter(f.suffix.lower() for f in files)
        for ext in SUPPORTED_EXTENSIONS:
            count = ext_counter.get(ext, 0)
            if count > 0:
                logger.info(f"  {ext}: {count}")

    def extract_from_pdf_streaming(self, file_path: Path) -> Generator[str, None, None]:
        """
        Потоковое извлечение текста из PDF для больших файлов

        Yields:
            Текст по частям
        """
        try:
            doc = pymupdf.open(file_path)
            total_pages = len(doc)

            logger.info(f"Обработка PDF: {file_path.name} ({total_pages} страниц)")

            # Обрабатываем батчами для экономии памяти
            for start_idx in range(0, total_pages, self.PAGE_BATCH_SIZE):
                end_idx = min(start_idx + self.PAGE_BATCH_SIZE, total_pages)
                batch_text = []

                for page_num in range(start_idx, end_idx):
                    try:
                        page = doc[page_num]
                        page_text = page.get_text()

                        # Фильтруем пустые или слишком короткие страницы
                        if len(page_text.strip()) >= self.MIN_PAGE_TEXT_LENGTH:
                            batch_text.append(page_text)

                    except Exception as e:
                        logger.warning(f"Ошибка при обработке страницы {page_num}: {e}")
                        continue

                if batch_text:
                    yield "\n".join(batch_text)

            doc.close()

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из PDF {file_path}: {e}")
            yield ""

    def extract_from_pdf(self, file_path: Path) -> str:
        """
        Извлечение текста из PDF с оптимизацией для больших файлов
        """
        try:
            doc = pymupdf.open(file_path)
            total_pages = len(doc)

            # Для очень больших PDF используем потоковую обработку
            if total_pages > self.MAX_PDF_PAGES_IN_MEMORY:
                logger.warning(
                    f"PDF содержит {total_pages} страниц. Используется потоковая обработка"
                )
                doc.close()

                # Собираем текст по частям
                text_parts = []
                for chunk in self.extract_from_pdf_streaming(file_path):
                    text_parts.append(chunk)

                return "\n".join(text_parts).strip()

            # Для обычных PDF - стандартная обработка
            text_parts = []

            for page_num in range(total_pages):
                try:
                    page = doc[page_num]
                    page_text = page.get_text()

                    if len(page_text.strip()) >= self.MIN_PAGE_TEXT_LENGTH:
                        text_parts.append(page_text)

                except Exception as e:
                    logger.warning(f"Ошибка при обработке страницы {page_num}: {e}")
                    continue

            doc.close()

            full_text = "\n".join(text_parts)
            logger.info(f"Извлечено {len(full_text)} символов из {total_pages} страниц")

            return full_text.strip()

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из PDF {file_path}: {e}")
            return ""

    def extract_from_docx(self, file_path: Path) -> str:
        """Извлечение текста из DOCX файла"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if (
                            cell_text and cell_text not in text_parts
                        ):  # Избегаем дубликатов
                            text_parts.append(cell_text)

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOCX {file_path}: {e}")
            return ""

    def extract_from_doc(self, file_path: Path) -> str:
        """Извлечение текста из DOC файла (только Windows)"""
        if not DOC_SUPPORT or self.word_app is None:
            logger.warning(f"Поддержка .doc файлов недоступна: {file_path}")
            return ""

        try:
            doc = self.word_app.Documents.Open(str(file_path.absolute()))
            text = doc.Content.Text
            doc.Close()
            return text.strip()

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOC {file_path}: {e}")
            # Попытка переинициализировать Word при ошибке
            self._init_word_app()
            return ""

    def extract_text(self, file_path: Path) -> str:
        """
        Универсальная функция извлечения текста

        Args:
            file_path: Путь к файлу

        Returns:
            Извлеченный текст
        """
        if not file_path.exists():
            logger.error(f"Файл не существует: {file_path}")
            return ""

        extension = file_path.suffix.lower()

        extractors = {
            ".pdf": self.extract_from_pdf,
            ".docx": self.extract_from_docx,
            ".doc": self.extract_from_doc,
        }

        extractor = extractors.get(extension)
        if extractor:
            return extractor(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат файла: {file_path}")
            return ""
